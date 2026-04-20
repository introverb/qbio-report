"""
One-off: convert MEDIA-DIRECTOR-HANDOFF.md to a polished .docx.

Run: python build_handoff_docx.py
Output: ../QUBIE-News-Media-Director-Handoff.docx
"""
import os
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Inches

HERE = os.path.dirname(os.path.abspath(__file__))
SRC  = os.path.join(HERE, "..", "MEDIA-DIRECTOR-HANDOFF.md")
OUT  = os.path.join(HERE, "..", "QUBIE-News-Media-Director-Handoff.docx")

# Brand palette (hex without #)
DARK_PURPLE = "2D1B30"
LIGHT_PURPLE = "7D4A6E"
PINK         = "D57DB2"
SOFT_TEXT    = "3A3036"
CODE_BG      = "F2F0EC"
GRAY_LINE    = "CCCCCC"
CREAM_DIM    = "E5DED2"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, color="CCCCCC", sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_hyperlink(paragraph, url, text, color=DARK_PURPLE):
    """Insert a clickable hyperlink into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # Color
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    # Underline
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ---- Inline parser --------------------------------------------------------
# Handles **bold**, *italic* / _italic_, `code`, and [text](url) inline.
_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*)"          # bold
    r"|(\*[^*]+\*)"              # italic with *
    r"|(_[^_]+_)"                # italic with _
    r"|(`[^`]+`)"                # inline code
    r"|(\[[^\]]+\]\([^)]+\))",   # [text](url)
)


def add_inline(paragraph, text, base_font="Calibri", base_size=11):
    """Parse a line for inline markdown and add runs to the paragraph."""
    # Replace em-dash-y sequences with unicode em dash for aesthetics
    text = text.replace(" -- ", " \u2014 ")
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            _add_plain_run(paragraph, text[pos:m.start()], base_font, base_size)
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            _add_plain_run(paragraph, token[2:-2], base_font, base_size, bold=True)
        elif (token.startswith("*") and token.endswith("*")) or \
             (token.startswith("_") and token.endswith("_")):
            _add_plain_run(paragraph, token[1:-1], base_font, base_size, italic=True)
        elif token.startswith("`") and token.endswith("`"):
            _add_code_run(paragraph, token[1:-1])
        elif token.startswith("[") and ")" in token:
            # [text](url)
            mlink = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if mlink:
                link_text, url = mlink.group(1), mlink.group(2)
                add_hyperlink(paragraph, url, link_text)
            else:
                _add_plain_run(paragraph, token, base_font, base_size)
        pos = m.end()
    if pos < len(text):
        _add_plain_run(paragraph, text[pos:], base_font, base_size)


def _add_plain_run(paragraph, text, font, size, bold=False, italic=False):
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    if bold:   run.bold = True
    if italic: run.italic = True


def _add_code_run(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    # light gray shading on the run
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CODE_BG)
    rPr.append(shd)


# ---- Main conversion ------------------------------------------------------

def convert():
    with open(SRC, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Page margins: 1 inch all around
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title header (top of the first page)
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run("QUBIE NEWS \u2014 Media Director Handoff")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(LIGHT_PURPLE)
    run.italic = True

    # Walk the markdown
    i = 0
    while i < len(lines):
        line = lines[i]

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), GRAY_LINE)
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph()
            if level == 1:
                run = p.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(22)
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(DARK_PURPLE)
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after  = Pt(8)
            elif level == 2:
                run = p.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(16)
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(DARK_PURPLE)
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after  = Pt(6)
            elif level == 3:
                run = p.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(13)
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(LIGHT_PURPLE)
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after  = Pt(4)
            else:
                run = p.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                run.bold = True
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        # Fenced code block
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            # Light-gray paragraph-level shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), CODE_BG)
            pPr.append(shd)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(8)
            continue

        # Table (pipe table). Detect by current line having pipes and next line being separator.
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", lines[i + 1]):
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            body_rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                body_rows.append(row)
                i += 1

            n_cols = max(len(header_cells), max((len(r) for r in body_rows), default=0))
            table = doc.add_table(rows=1 + len(body_rows), cols=n_cols)
            table.autofit = True

            # Header row
            for c_i in range(n_cols):
                cell = table.rows[0].cells[c_i]
                cell.text = ""
                set_cell_shading(cell, DARK_PURPLE)
                set_cell_border(cell)
                para = cell.paragraphs[0]
                txt = header_cells[c_i] if c_i < len(header_cells) else ""
                # Header cells: white bold
                run = para.add_run(txt)
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor.from_string("FFFFFF")

            # Body rows
            for r_i, row in enumerate(body_rows, start=1):
                for c_i in range(n_cols):
                    cell = table.rows[r_i].cells[c_i]
                    cell.text = ""
                    set_cell_border(cell)
                    # Alternating shading for readability
                    if r_i % 2 == 0:
                        set_cell_shading(cell, CREAM_DIM)
                    para = cell.paragraphs[0]
                    text = row[c_i] if c_i < len(row) else ""
                    add_inline(para, text)

            # small spacing after the table
            doc.add_paragraph()
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                m2 = re.match(r"^(\s*)[-*]\s+(.*)$", lines[i])
                indent_spaces = len(m2.group(1))
                text = m2.group(2)
                # Map indent: 2 spaces = 1 level
                level = indent_spaces // 2
                p = doc.add_paragraph(style="List Bullet")
                if level:
                    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
                add_inline(p, text)
                i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                m2 = re.match(r"^\s*\d+\.\s+(.*)$", lines[i])
                text = m2.group(1)
                p = doc.add_paragraph(style="List Number")
                add_inline(p, text)
                i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_inline(p, line)
        i += 1

    # Remove the blank first paragraph Word sometimes auto-inserts (if present, our title header is first)
    doc.save(os.path.abspath(OUT))
    print(f"Saved: {os.path.abspath(OUT)}")


if __name__ == "__main__":
    convert()
